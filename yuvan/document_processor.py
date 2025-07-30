"""
Document Processing System for Yuvan AI Assistant
Handles PDF, TXT, DOCX files with OCR capabilities using Tesseract and LangChain loaders
"""

import os
import io
import tempfile
from typing import Dict, List, Any, Optional, Union
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
import hashlib

# Document processing libraries
from langchain.document_loaders import (
    PyPDFLoader, 
    TextLoader, 
    Docx2txtLoader,
    UnstructuredPDFLoader
)
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document

# OCR libraries
try:
    import pytesseract
    from PIL import Image
    import fitz  # PyMuPDF for PDF image extraction
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# Additional libraries
import magic
import pandas as pd
import json
import requests
from urllib.parse import urlparse

@dataclass
class ProcessedDocument:
    """Represents a processed document"""
    file_path: str
    file_type: str
    file_size: int
    content: str
    metadata: Dict[str, Any]
    chunks: List[Document]
    summary: str
    timestamp: datetime
    processing_time: float
    ocr_used: bool = False
    confidence_score: float = 1.0

class DocumentProcessor:
    """Advanced document processing with OCR and LangChain integration"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        # Supported file types
        self.supported_types = {
            '.pdf': self._process_pdf,
            '.txt': self._process_txt,
            '.docx': self._process_docx,
            '.doc': self._process_doc,
            '.rtf': self._process_rtf,
            '.md': self._process_markdown,
            '.html': self._process_html,
            '.csv': self._process_csv,
            '.json': self._process_json
        }
        
        # OCR configuration
        if OCR_AVAILABLE:
            # Configure Tesseract (adjust path as needed)
            try:
                # Common Tesseract paths
                tesseract_paths = [
                    '/usr/bin/tesseract',
                    '/opt/homebrew/bin/tesseract',
                    r'C:\Program Files\Tesseract-OCR\tesseract.exe',
                    r'C:\Users\User\AppData\Local\Programs\Tesseract-OCR\tesseract.exe'
                ]
                
                for path in tesseract_paths:
                    if os.path.exists(path):
                        pytesseract.pytesseract.tesseract_cmd = path
                        break
                        
            except Exception as e:
                print(f"Warning: Tesseract configuration issue: {e}")
        
        # Document cache
        self.processed_documents: Dict[str, ProcessedDocument] = {}
    
    def process_document(self, file_path: str, use_ocr: bool = True, 
                        summarize: bool = True) -> ProcessedDocument:
        """Process a document and extract content"""
        start_time = datetime.now()
        
        # Check if file exists
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Get file information
        file_info = self._get_file_info(file_path)
        
        # Check cache
        file_hash = self._get_file_hash(file_path)
        if file_hash in self.processed_documents:
            return self.processed_documents[file_hash]
        
        # Determine file type and process
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext not in self.supported_types:
            raise ValueError(f"Unsupported file type: {file_ext}")
        
        try:
            # Process the document
            processor = self.supported_types[file_ext]
            content, metadata, ocr_used = processor(file_path, use_ocr)
            
            # Create chunks
            chunks = self._create_chunks(content, file_path)
            
            # Generate summary if requested
            summary = ""
            if summarize and content:
                summary = self._generate_summary(content)
            
            # Calculate processing time
            processing_time = (datetime.now() - start_time).total_seconds()
            
            # Create processed document
            processed_doc = ProcessedDocument(
                file_path=file_path,
                file_type=file_ext,
                file_size=file_info['size'],
                content=content,
                metadata={**file_info, **metadata},
                chunks=chunks,
                summary=summary,
                timestamp=datetime.now(),
                processing_time=processing_time,
                ocr_used=ocr_used,
                confidence_score=metadata.get('confidence', 1.0)
            )
            
            # Cache the result
            self.processed_documents[file_hash] = processed_doc
            
            return processed_doc
            
        except Exception as e:
            raise Exception(f"Error processing document: {str(e)}")
    
    def _get_file_info(self, file_path: str) -> Dict[str, Any]:
        """Get basic file information"""
        file_stats = os.stat(file_path)
        
        # Try to detect MIME type
        mime_type = "unknown"
        try:
            mime_type = magic.from_file(file_path, mime=True)
        except Exception:
            pass
        
        return {
            'filename': os.path.basename(file_path),
            'size': file_stats.st_size,
            'created': datetime.fromtimestamp(file_stats.st_ctime),
            'modified': datetime.fromtimestamp(file_stats.st_mtime),
            'mime_type': mime_type,
            'extension': Path(file_path).suffix.lower()
        }
    
    def _get_file_hash(self, file_path: str) -> str:
        """Generate hash for file caching"""
        with open(file_path, 'rb') as f:
            file_hash = hashlib.md5(f.read()).hexdigest()
        return file_hash
    
    def _process_pdf(self, file_path: str, use_ocr: bool = True) -> tuple:
        """Process PDF file with optional OCR"""
        content = ""
        metadata = {}
        ocr_used = False
        
        try:
            # First try standard PDF text extraction
            loader = PyPDFLoader(file_path)
            documents = loader.load()
            
            # Combine text from all pages
            text_content = "\n".join([doc.page_content for doc in documents])
            
            # If text extraction yields little content, try OCR
            if use_ocr and OCR_AVAILABLE and len(text_content.strip()) < 100:
                ocr_content = self._extract_text_with_ocr(file_path)
                if len(ocr_content) > len(text_content):
                    text_content = ocr_content
                    ocr_used = True
            
            content = text_content
            metadata = {
                'pages': len(documents),
                'extraction_method': 'OCR' if ocr_used else 'text',
                'confidence': 0.8 if ocr_used else 1.0
            }
            
            # Try to extract additional PDF metadata
            try:
                with fitz.open(file_path) as pdf_doc:
                    pdf_metadata = pdf_doc.metadata
                    metadata.update({
                        'title': pdf_metadata.get('title', ''),
                        'author': pdf_metadata.get('author', ''),
                        'subject': pdf_metadata.get('subject', ''),
                        'creator': pdf_metadata.get('creator', ''),
                        'producer': pdf_metadata.get('producer', ''),
                        'creation_date': pdf_metadata.get('creationDate', ''),
                        'modification_date': pdf_metadata.get('modDate', '')
                    })
            except Exception:
                pass
                
        except Exception as e:
            # Fallback to unstructured loader
            try:
                loader = UnstructuredPDFLoader(file_path)
                documents = loader.load()
                content = "\n".join([doc.page_content for doc in documents])
                metadata = {'extraction_method': 'unstructured', 'confidence': 0.9}
            except Exception:
                if use_ocr and OCR_AVAILABLE:
                    content = self._extract_text_with_ocr(file_path)
                    ocr_used = True
                    metadata = {'extraction_method': 'OCR_only', 'confidence': 0.7}
                else:
                    raise Exception(f"Failed to process PDF: {str(e)}")
        
        return content, metadata, ocr_used
    
    def _process_txt(self, file_path: str, use_ocr: bool = False) -> tuple:
        """Process text file"""
        try:
            loader = TextLoader(file_path, encoding='utf-8')
            documents = loader.load()
            content = documents[0].page_content if documents else ""
            
            metadata = {
                'encoding': 'utf-8',
                'extraction_method': 'direct',
                'confidence': 1.0
            }
            
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin1', 'cp1252', 'iso-8859-1']:
                try:
                    loader = TextLoader(file_path, encoding=encoding)
                    documents = loader.load()
                    content = documents[0].page_content if documents else ""
                    metadata = {
                        'encoding': encoding,
                        'extraction_method': 'direct',
                        'confidence': 0.9
                    }
                    break
                except Exception:
                    continue
            else:
                raise Exception("Could not decode text file with any encoding")
        
        return content, metadata, False
    
    def _process_docx(self, file_path: str, use_ocr: bool = False) -> tuple:
        """Process DOCX file"""
        try:
            loader = Docx2txtLoader(file_path)
            documents = loader.load()
            content = documents[0].page_content if documents else ""
            
            metadata = {
                'extraction_method': 'docx2txt',
                'confidence': 1.0
            }
            
            # Try to extract additional metadata using python-docx
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(file_path)
                
                core_props = doc.core_properties
                metadata.update({
                    'title': core_props.title or '',
                    'author': core_props.author or '',
                    'subject': core_props.subject or '',
                    'created': core_props.created,
                    'modified': core_props.modified,
                    'last_modified_by': core_props.last_modified_by or '',
                    'paragraphs': len(doc.paragraphs)
                })
            except Exception:
                pass
                
        except Exception as e:
            raise Exception(f"Failed to process DOCX: {str(e)}")
        
        return content, metadata, False
    
    def _process_doc(self, file_path: str, use_ocr: bool = False) -> tuple:
        """Process legacy DOC file"""
        # This would require additional libraries like python-docx2txt or conversion
        raise NotImplementedError("Legacy DOC files require additional libraries")
    
    def _process_rtf(self, file_path: str, use_ocr: bool = False) -> tuple:
        """Process RTF file"""
        # This would require striprtf or similar library
        raise NotImplementedError("RTF files require additional libraries")
    
    def _process_markdown(self, file_path: str, use_ocr: bool = False) -> tuple:
        """Process Markdown file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        metadata = {
            'extraction_method': 'direct',
            'confidence': 1.0,
            'format': 'markdown'
        }
        
        return content, metadata, False
    
    def _process_html(self, file_path: str, use_ocr: bool = False) -> tuple:
        """Process HTML file"""
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Extract text content
            content = soup.get_text()
            
            # Extract metadata
            title = soup.find('title')
            meta_tags = soup.find_all('meta')
            
            metadata = {
                'extraction_method': 'beautifulsoup',
                'confidence': 1.0,
                'title': title.text if title else '',
                'format': 'html'
            }
            
            # Extract meta information
            for meta in meta_tags:
                name = meta.get('name', '')
                content_attr = meta.get('content', '')
                if name and content_attr:
                    metadata[f'meta_{name}'] = content_attr
                    
        except Exception as e:
            # Fallback to plain text
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            metadata = {'extraction_method': 'fallback', 'confidence': 0.7}
        
        return content, metadata, False
    
    def _process_csv(self, file_path: str, use_ocr: bool = False) -> tuple:
        """Process CSV file"""
        try:
            df = pd.read_csv(file_path)
            
            # Convert to text representation
            content = f"CSV Data Summary:\n"
            content += f"Rows: {len(df)}, Columns: {len(df.columns)}\n\n"
            content += f"Column Names: {', '.join(df.columns)}\n\n"
            content += f"First 5 rows:\n{df.head().to_string()}\n\n"
            content += f"Data types:\n{df.dtypes.to_string()}\n\n"
            
            if len(df) <= 100:  # For small files, include all data
                content += f"Complete data:\n{df.to_string()}"
            
            metadata = {
                'extraction_method': 'pandas',
                'confidence': 1.0,
                'rows': len(df),
                'columns': len(df.columns),
                'column_names': list(df.columns),
                'data_types': df.dtypes.to_dict()
            }
            
        except Exception as e:
            raise Exception(f"Failed to process CSV: {str(e)}")
        
        return content, metadata, False
    
    def _process_json(self, file_path: str, use_ocr: bool = False) -> tuple:
        """Process JSON file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Convert to readable text
            content = f"JSON Data Structure:\n"
            content += json.dumps(data, indent=2, ensure_ascii=False)
            
            metadata = {
                'extraction_method': 'json',
                'confidence': 1.0,
                'data_type': type(data).__name__,
                'keys': list(data.keys()) if isinstance(data, dict) else None
            }
            
        except Exception as e:
            raise Exception(f"Failed to process JSON: {str(e)}")
        
        return content, metadata, False
    
    def _extract_text_with_ocr(self, file_path: str) -> str:
        """Extract text using OCR (Tesseract)"""
        if not OCR_AVAILABLE:
            raise Exception("OCR libraries not available")
        
        try:
            extracted_text = ""
            
            # Open PDF and extract images
            pdf_document = fitz.open(file_path)
            
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                
                # Get images on the page
                image_list = page.get_images()
                
                if image_list:
                    # Process images with OCR
                    for img_index, img in enumerate(image_list):
                        xref = img[0]
                        pix = fitz.Pixmap(pdf_document, xref)
                        
                        if pix.n - pix.alpha < 4:  # GRAY or RGB
                            img_data = pix.tobytes("png")
                            image = Image.open(io.BytesIO(img_data))
                            
                            # Apply OCR
                            text = pytesseract.image_to_string(
                                image, 
                                config='--psm 6 --oem 3'
                            )
                            extracted_text += text + "\n"
                        
                        pix = None
                else:
                    # If no images, render page as image and OCR
                    mat = fitz.Matrix(2, 2)  # Increase resolution
                    pix = page.get_pixmap(matrix=mat)
                    img_data = pix.tobytes("png")
                    image = Image.open(io.BytesIO(img_data))
                    
                    text = pytesseract.image_to_string(
                        image,
                        config='--psm 6 --oem 3'
                    )
                    extracted_text += text + "\n"
                    
                    pix = None
            
            pdf_document.close()
            return extracted_text.strip()
            
        except Exception as e:
            raise Exception(f"OCR extraction failed: {str(e)}")
    
    def _create_chunks(self, content: str, source: str) -> List[Document]:
        """Create document chunks for better processing"""
        if not content:
            return []
        
        # Split the text into chunks
        texts = self.text_splitter.split_text(content)
        
        # Create Document objects
        chunks = []
        for i, text in enumerate(texts):
            chunk = Document(
                page_content=text,
                metadata={
                    'source': source,
                    'chunk_index': i,
                    'chunk_size': len(text)
                }
            )
            chunks.append(chunk)
        
        return chunks
    
    def _generate_summary(self, content: str, max_length: int = 500) -> str:
        """Generate a simple extractive summary"""
        if len(content) <= max_length:
            return content
        
        # Simple extractive summary - take first few sentences
        sentences = content.split('. ')
        summary = ""
        
        for sentence in sentences:
            if len(summary + sentence) < max_length:
                summary += sentence + ". "
            else:
                break
        
        if not summary:
            summary = content[:max_length] + "..."
        
        return summary.strip()
    
    def process_url(self, url: str, use_ocr: bool = False) -> ProcessedDocument:
        """Download and process a document from URL"""
        try:
            # Download the file
            response = requests.get(url, timeout=30)
            response.raise_for_status()
            
            # Determine file type from URL or content type
            parsed_url = urlparse(url)
            file_ext = Path(parsed_url.path).suffix.lower()
            
            if not file_ext:
                content_type = response.headers.get('content-type', '')
                if 'pdf' in content_type:
                    file_ext = '.pdf'
                elif 'text' in content_type:
                    file_ext = '.txt'
                else:
                    file_ext = '.txt'  # Default
            
            # Save to temporary file
            with tempfile.NamedTemporaryFile(suffix=file_ext, delete=False) as temp_file:
                temp_file.write(response.content)
                temp_path = temp_file.name
            
            try:
                # Process the temporary file
                result = self.process_document(temp_path, use_ocr)
                result.file_path = url  # Update to show original URL
                return result
            finally:
                # Clean up temporary file
                os.unlink(temp_path)
                
        except Exception as e:
            raise Exception(f"Failed to process URL: {str(e)}")
    
    def get_document_stats(self) -> Dict[str, Any]:
        """Get statistics about processed documents"""
        if not self.processed_documents:
            return {"total_documents": 0}
        
        stats = {
            "total_documents": len(self.processed_documents),
            "file_types": {},
            "total_size": 0,
            "ocr_usage": 0,
            "average_processing_time": 0
        }
        
        total_time = 0
        for doc in self.processed_documents.values():
            # File types
            file_type = doc.file_type
            stats["file_types"][file_type] = stats["file_types"].get(file_type, 0) + 1
            
            # Total size
            stats["total_size"] += doc.file_size
            
            # OCR usage
            if doc.ocr_used:
                stats["ocr_usage"] += 1
            
            # Processing time
            total_time += doc.processing_time
        
        stats["average_processing_time"] = total_time / len(self.processed_documents)
        
        return stats
    
    def search_documents(self, query: str, max_results: int = 10) -> List[ProcessedDocument]:
        """Search processed documents for a query"""
        results = []
        query_lower = query.lower()
        
        for doc in self.processed_documents.values():
            if query_lower in doc.content.lower():
                results.append(doc)
        
        # Sort by relevance (simple word count for now)
        results.sort(key=lambda d: d.content.lower().count(query_lower), reverse=True)
        
        return results[:max_results]
    
    def clear_cache(self):
        """Clear processed documents cache"""
        self.processed_documents.clear()